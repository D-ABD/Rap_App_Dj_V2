# rap_app/utils/exporter.py

import io
import csv

from django.http import HttpResponse
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4


class Exporter:
    """
    🧰 Classe utilitaire pour exporter un queryset Django dans plusieurs formats :
    - CSV
    - Word (.docx)
    - PDF (via ReportLab)

    Usage :
        exporter = Exporter(queryset, fields=["nom", "created_by.username"], headers=["Nom", "Auteur"])
        return exporter.export_csv("export.csv")
    """

    def __init__(self, queryset, fields, headers=None):
        """
        :param queryset: QuerySet Django à exporter
        :param fields: Liste des champs à extraire (ex: "formation.nom", "created_by.username")
        :param headers: Noms de colonnes (facultatif – sinon mêmes que fields)
        """
        self.queryset = queryset
        self.fields = fields
        self.headers = headers or fields

    def _resolve_field(self, obj, field_path):
        """
        Résout dynamiquement les relations sur les objets (ex: "formation.nom")
        """
        for part in field_path.split("."):
            obj = getattr(obj, part, None)
            if callable(obj):
                obj = obj()
        return str(obj) if obj is not None else ""

    def get_data(self):
        """
        Génère les données ligne par ligne à partir du queryset
        """
        for obj in self.queryset:
            yield [self._resolve_field(obj, f) for f in self.fields]

    def export_csv(self, filename="export.csv"):
        """
        Exporte les données en CSV
        """
        buffer = io.StringIO()
        writer = csv.writer(buffer)
        writer.writerow(self.headers)
        for row in self.get_data():
            writer.writerow(row)
        buffer.seek(0)
        return HttpResponse(
            buffer.getvalue(),
            content_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    def export_word(self, filename="export.docx"):
        """
        Exporte les données en document Word (.docx)
        """
        doc = Document()
        doc.add_heading("Export de données", 0)

        table = doc.add_table(rows=1, cols=len(self.headers))
        for i, h in enumerate(self.headers):
            table.cell(0, i).text = h

        for row in self.get_data():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = val

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return HttpResponse(
            f.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    def export_pdf(self, filename="export.pdf"):
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        margin = 40
        y = height - margin

        def draw_header():
            p.setFont("Helvetica-Bold", 14)
            p.drawString(margin, y, "Export de données")
            return y - 30

        def draw_row(row_data):
            nonlocal y
            p.setFont("Helvetica", 10)
            for i, val in enumerate(row_data):
                # Découpage du texte long
                lines = split_text(val, max_length=90)
                for line in lines:
                    if y < 50:
                        p.showPage()
                        y = height - margin
                    p.drawString(margin, y, f"{self.headers[i]} : {line}")
                    y -= 15
            y -= 10  # espacement entre objets

        def split_text(text, max_length=90):
            return [text[i:i+max_length] for i in range(0, len(text), max_length)]

        y = draw_header()
        for row in self.get_data():
            if y < 100:
                p.showPage()
                y = height - margin
                y = draw_header()
            draw_row(row)

        p.save()
        buffer.seek(0)
        return HttpResponse(
            buffer.read(),
            content_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Length": str(buffer.getbuffer().nbytes),
            }
        )
